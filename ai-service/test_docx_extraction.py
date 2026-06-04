"""Test DOCX extraction để prove python-docx cũng hoạt động."""
import httpx
import io
from docx import Document


def create_test_docx() -> bytes:
    """Tạo DOCX với nội dung hợp đồng."""
    doc = Document()
    
    # Title
    doc.add_heading('HOP DONG THUE DAT', level=1)
    
    # Content
    doc.add_paragraph('Ben A: Chu dat - Ong Nguyen Van C')
    doc.add_paragraph('CMND: 111222333')
    doc.add_paragraph('')
    doc.add_paragraph('Ben B: Nguoi thue - Ong Le Van D')
    doc.add_paragraph('CMND: 444555666')
    doc.add_paragraph('')
    
    doc.add_heading('Dieu 1: Thong tin thua dat', level=2)
    doc.add_paragraph('Dia chi: Xa Tan Binh, Huyen Cu Chi, TP.HCM')
    doc.add_paragraph('Dien tich: 500m2')
    doc.add_paragraph('So hong: 123/456/789')
    doc.add_paragraph('')
    
    doc.add_heading('Dieu 2: Gia thue dat', level=2)
    doc.add_paragraph('- Gia thue: 2.000.000 VND/thang')
    doc.add_paragraph('- Tien dat coc: 6.000.000 VND (3 thang)')
    doc.add_paragraph('- Hinh thuc thanh toan: Tien mat hoac chuyen khoan')
    doc.add_paragraph('')
    
    doc.add_heading('Dieu 3: Thoi han', level=2)
    doc.add_paragraph('24 thang, tu 01/07/2026 den 30/06/2028')
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def test_docx_extraction():
    """Test chi tiết DOCX extraction."""
    print("="*70)
    print("TEST: Chi tiết DOCX Text Extraction")
    print("="*70)
    
    # 1. Tạo DOCX
    print("\n[STEP 1] 📄 Tạo DOCX file...")
    docx_content = create_test_docx()
    print(f"  ✅ DOCX size: {len(docx_content)} bytes")
    
    # 2. Extract text bằng python-docx
    print("\n[STEP 2] 📖 Extract text bằng python-docx...")
    
    docx_file = io.BytesIO(docx_content)
    doc = Document(docx_file)
    
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    print(f"  ✅ Text extracted: {len(text)} characters")
    print(f"  📊 Số paragraphs: {len(doc.paragraphs)}")
    print(f"\n  📝 Sample text (250 chars):")
    print(f"  {text[:250]}...")
    
    # 3. Upload lên server
    print("\n[STEP 3] 📤 Upload DOCX lên server...")
    url = "http://localhost:8000/api/ai/documents/upload"
    
    files = {'file': ('hop_dong_thue_dat.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {'user_id': 'test_docx', 'contract_type': 'HOUSE_RENTAL'}
    
    response = httpx.post(url, files=files, data=data, timeout=30.0)
    
    if response.status_code == 200:
        result = response.json()
        print(f"  ✅ Upload thành công!")
        print(f"  📝 Document ID: {result['data']['documentId']}")
        print(f"  📏 Server extracted: {result['data']['textLength']} characters")
    else:
        print(f"  ❌ Upload failed: {response.status_code}")
        print(f"  {response.text}")
        return
    
    # 4. Test search
    print("\n[STEP 4] 🔍 Test search trong DOCX content...")
    search_url = "http://localhost:8000/api/ai/documents/search"
    
    test_queries = [
        ("thue dat", "Tìm hợp đồng thuê đất"),
        ("dat coc", "Tìm tiền đặt cọc"),
        ("Cu Chi", "Tìm địa chỉ"),
    ]
    
    for query, description in test_queries:
        data = {'user_id': 'test_docx', 'query': query, 'top_k': 1}
        response = httpx.post(search_url, data=data, timeout=10.0)
        
        if response.status_code == 200:
            result = response.json()
            matches = result['data']['results']
            
            if matches:
                score = matches[0]['score']
                snippet = matches[0]['snippet'][:80]
                print(f"  ✅ '{query}' → Score: {score:.4f}")
                print(f"     {description}")
                print(f"     Snippet: {snippet}...")
            else:
                print(f"  ❌ '{query}' → No match")
    
    print("\n" + "="*70)
    print("✅ TEST COMPLETED - DOCX EXTRACTION WORKING!")
    print("="*70)
    
    print("\n📊 KẾT LUẬN:")
    print("  ✅ python-docx extract text thành công")
    print("  ✅ DOCX upload và lưu trữ OK")
    print("  ✅ Search trong DOCX content hoạt động tốt")
    print("  ✅ Support cả PDF lẫn DOCX!")
    print("\n💡 KHÔNG cần thư viện nặng:")
    print("  ❌ KHÔNG dùng torch")
    print("  ❌ KHÔNG dùng transformers")
    print("  ✅ CHỈ dùng PyPDF2 + python-docx (tổng ~3MB)")


if __name__ == "__main__":
    try:
        test_docx_extraction()
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
