from __future__ import annotations

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Iterable

from app.models.knowledge_models import DocumentBlock, ExtractedDocument


LEGAL_SECTION_PATTERNS = (
    re.compile(r"^\s*Chương\s+[IVXLC\d]+(?:\s+.*)?$", re.IGNORECASE),
    re.compile(r"^\s*Mục\s+\d+(?:\s+.*)?$", re.IGNORECASE),
    re.compile(r"^\s*Điều\s+\d+(?:[.\-:]\s*.*)?$", re.IGNORECASE),
    re.compile(r"^\s*Khoản\s+\d+(?:[.\-:]\s*.*)?$", re.IGNORECASE),
    re.compile(r"^\s*Điểm\s+[a-zA-Z0-9]+(?:[.\-:]\s*.*)?$", re.IGNORECASE),
)

QUESTION_PATTERNS = (
    re.compile(r"^\s*(Q(?:uestion)?|Câu hỏi)\s*[:\-]", re.IGNORECASE),
    re.compile(r".+\?\s*$"),
)
def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _is_heading_like(text: str) -> bool:
    return any(pattern.match(text) for pattern in LEGAL_SECTION_PATTERNS)


def _looks_like_question(text: str) -> bool:
    return any(pattern.match(text) for pattern in QUESTION_PATTERNS)


def _group_lines_into_blocks(lines: Iterable[str]) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if current:
                blocks.append(_normalize_text(" ".join(current)))
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(_normalize_text(" ".join(current)))

    return [block for block in blocks if block]


class DocumentLoader(ABC):
    supported_extensions: tuple[str, ...] = ()

    @abstractmethod
    def load(self, source_path: Path) -> ExtractedDocument:
        raise NotImplementedError

    def supports(self, source_path: Path) -> bool:
        return source_path.suffix.lower() in self.supported_extensions


class PdfDocumentLoader(DocumentLoader):
    supported_extensions = (".pdf",)

    def load(self, source_path: Path) -> ExtractedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(source_path))
        blocks: list[DocumentBlock] = []

        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            page_blocks = _group_lines_into_blocks(page_text.splitlines())
            for block_index, block in enumerate(page_blocks, start=1):
                kind = "paragraph"
                heading_level = None
                if _is_heading_like(block):
                    kind = "heading"
                    heading_level = 1
                elif _looks_like_question(block):
                    kind = "question"
                blocks.append(
                    DocumentBlock(
                        text=block,
                        kind=kind,
                        order=len(blocks) + 1,
                        page_number=page_index,
                        heading_level=heading_level,
                        metadata={"page_block_index": block_index},
                    )
                )

        return ExtractedDocument(
            source_path=source_path,
            title=source_path.stem,
            file_type="pdf",
            blocks=blocks,
            metadata={"page_count": len(reader.pages)},
        )


class DocxDocumentLoader(DocumentLoader):
    supported_extensions = (".docx",)

    def load(self, source_path: Path) -> ExtractedDocument:
        from docx import Document

        doc = Document(str(source_path))
        blocks: list[DocumentBlock] = []

        for paragraph in doc.paragraphs:
            text = _normalize_text(paragraph.text)
            if not text:
                continue

            style_name = getattr(paragraph.style, "name", None)
            heading_level = None
            kind = "paragraph"
            if style_name:
                heading_match = re.match(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
                if heading_match:
                    heading_level = int(heading_match.group(1))
                    kind = "heading"
            if _is_heading_like(text):
                kind = "heading"
                heading_level = heading_level or 1
            elif _looks_like_question(text):
                kind = "question"

            blocks.append(
                DocumentBlock(
                    text=text,
                    kind=kind,
                    order=len(blocks) + 1,
                    style=style_name,
                    heading_level=heading_level,
                )
            )

        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                row_text = " | ".join(_normalize_text(cell.text) for cell in row.cells if _normalize_text(cell.text))
                if not row_text:
                    continue
                blocks.append(
                    DocumentBlock(
                        text=row_text,
                        kind="table_row",
                        order=len(blocks) + 1,
                        metadata={"table_index": table_index, "row_index": row_index},
                    )
                )

        return ExtractedDocument(
            source_path=source_path,
            title=source_path.stem,
            file_type="docx",
            blocks=blocks,
            metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
        )


class LoaderRegistry:
    def __init__(self) -> None:
        self._loaders: list[DocumentLoader] = []

    def register(self, loader: DocumentLoader) -> None:
        self._loaders.append(loader)

    def supported_extensions(self) -> list[str]:
        extensions: list[str] = []
        for loader in self._loaders:
            extensions.extend(loader.supported_extensions)
        return sorted(set(extensions))

    def resolve(self, source_path: Path) -> DocumentLoader:
        for loader in self._loaders:
            if loader.supports(source_path):
                return loader
        raise ValueError(f"Unsupported file type: {source_path.suffix.lower()}")


def build_default_loader_registry() -> LoaderRegistry:
    registry = LoaderRegistry()
    registry.register(PdfDocumentLoader())
    registry.register(DocxDocumentLoader())
    return registry

