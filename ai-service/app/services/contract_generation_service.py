from __future__ import annotations

import logging
import os
import re
import shutil
import uuid

import httpx
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import HTTPException

from app.schemas import RagCitation, RagQueryRequest, RagQueryResponse, RagUsage
from app.services.llm_client import build_default_llm_client
from app.services.retrieval_service import RagChunkHit, RetrievalService

logger = logging.getLogger(__name__)


def is_contract_generation_intent(question: str) -> bool:
    """Detect whether the user is asking for a contract/template file download.

    Uses fast keyword matching — no LLM call needed.
    """
    question_lower = question.lower().strip()

    # Informational or QA keywords that should bypass contract generation
    qa_keywords = [
        "what is", "why ", "how to", "how do", "how can", "is there", "are there",
        "là gì", "như thế nào", "làm sao", "làm thế nào", "có được", "có nên",
        "tại sao", "giải thích", "khái niệm", "quy định", "luật nào", "thủ tục",
        "cần lưu ý gì", "rủi ro gì", "bắt buộc", "yêu cầu", "điều kiện",
        "liệt kê", "danh sách", "kể tên", "cho biết các", "những hợp đồng nào",
        "gồm những", "có bao nhiêu", "ở đâu", "từ đâu", "phân tích", "so sánh",
        "kiểm tra", "đánh giá", "nhận xét", "góp ý", "sửa lỗi", "rà soát",
        "nguồn nào", "file nào", "tệp nào", "thông tin lấy từ", "lấy từ đâu",
        "tham khảo", "đã tham khảo", "tham chiếu", "đã dùng", "đã sử dụng", "nguồn gốc",
        "căn cứ", "căn cứ vào", "dựa vào", "dựa trên", "căn cứ đâu", "bạn lấy ở đâu",
        "dẫn chứng", "bằng chứng", "nghị định nào", "thông tư nào", "luật nào",
    ]
    if any(kw in question_lower for kw in qa_keywords):
        return False

    creation_patterns = [
        r"\b(?:generate|create|draft|provide|give|make|write)\b",
        r"(?<!\w)(?:tạo|soạn thảo|soạn|viết|cung cấp|lập|in|xuất|tải|lấy|mẫu)(?!\w)",
        r"(?<!\w)cho tôi(?!\w)",
        r"(?<!\w)bản thảo(?!\w)",
    ]
    contract_keywords = [
        "contract", "agreement", "lease", "tenancy", "rental",
        "hợp đồng", "thoả thuận", "thuê nhà", "thuê trọ", "thuê phòng", "thuê văn phòng", "thuê mặt bằng",
        "lao động", "làm thêm", "thực tập", "cộng tác viên", "freelance", "mua bán tài sản cá nhân", "vay tiền cá nhân",
        "mua bán", "dịch vụ", "chuyển nhượng", "tặng cho", "vay tiền", "ủy quyền",
    ]

    has_creation = any(re.search(pattern, question_lower) for pattern in creation_patterns)
    has_contract = any(kw in question_lower for kw in contract_keywords)

    direct_patterns = [
        "draft a contract", "draft an agreement", "create a contract", "create an agreement",
        "generate a contract", "generate an agreement", "write a contract", "write an agreement",
        "mẫu hợp đồng", "soạn hợp đồng", "soạn thảo hợp đồng", "tạo hợp đồng", "viết hợp đồng",
    ]
    has_direct = any(pat in question_lower for pat in direct_patterns)

    return (has_creation and has_contract) or has_direct


def is_export_docx_intent(question: str) -> bool:
    """Detect whether the user is asking to export/download the current chat content as a DOCX file.
    """
    question_lower = question.lower().strip()

    # Informational reference document queries should NOT be treated as file export requests
    reference_keywords = [
        "tham khảo", "tài liệu tham khảo", "văn bản tham khảo", "tài liệu khác",
        "danh sách tài liệu", "danh sách văn bản", "nguồn tham khảo", "căn cứ pháp lý",
    ]
    if any(kw in question_lower for kw in reference_keywords) and not any(kw in question_lower for kw in ["docx", "doc", "word", "pdf"]):
        return False

    # ── 1. Export/action keywords ──
    export_keywords = [
        # Vietnamese
        "xuất", "tải về", "tải xuống", "tải cho", "tải file",
        "lưu", "lưu lại", "lưu về", "lưu thành",
        "gửi", "gửi cho", "gửi file", "gửi tệp",
        "in ra", "in file", "in thành",
        "chuyển", "chuyển sang", "chuyển thành", "chuyển đổi",
        "tạo file", "tạo tệp", "làm file",
        "cho tôi", "đưa cho tôi", "gửi cho tôi", "gửi tôi",
        "copy ra", "sao chép ra",
        # English
        "export", "download", "save", "print",
        "convert", "generate file", "create file",
        "give me", "send me", "get file",
    ]

    # ── 2. File format keywords ──
    file_keywords = [
        # Specific formats
        "docx", "doc", "word", ".docx", ".doc",
        "pdf", ".pdf",
        # Generic file terms - Vietnamese
        "file", "tệp", "tập tin", "tệp tin",
        "văn bản", "tài liệu", "bản word", "bản doc",
        # Generic file terms - English
        "document",
    ]

    has_export = any(kw in question_lower for kw in export_keywords)
    has_file = any(kw in question_lower for kw in file_keywords)

    # ── 3. Direct phrase patterns (high confidence) ──
    direct_patterns = [
        # === XUẤT (export) ===
        "xuất file", "xuất ra file", "xuất thành file",
        "xuất docx", "xuất doc", "xuất word", "xuất pdf",
        "xuất ra docx", "xuất ra doc", "xuất ra word", "xuất ra pdf",
        "xuất cho tôi", "xuất giúp tôi", "xuất giùm tôi",
        "xuất cho tôi file", "xuất cho tôi ra file",
        "xuất file cho tôi", "xuất cho tôi ra",
        "xuất bản nháp", "xuất hợp đồng", "xuất nội dung",
        "xuất ra bản", "xuất thành bản",

        # === TẢI (download) ===
        "tải file", "tải docx", "tải doc", "tải word", "tải pdf",
        "tải về file", "tải về docx", "tải về word",
        "tải xuống file", "tải xuống docx",
        "tải cho tôi", "tải giúp tôi", "tải giùm tôi",
        "download file", "download docx", "download word", "download pdf",

        # === LƯU (save) ===
        "lưu file", "lưu docx", "lưu doc", "lưu word",
        "lưu thành file", "lưu thành docx", "lưu thành word",
        "lưu ra file", "lưu lại file", "lưu về file",
        "save file", "save docx", "save as docx", "save as word",
        "save to file", "save to docx",

        # === IN (print) ===
        "in ra file", "in ra docx", "in ra word",
        "in thành file", "in file",

        # === CHUYỂN (convert) ===
        "chuyển sang docx", "chuyển sang word", "chuyển sang file",
        "chuyển thành file", "chuyển thành docx", "chuyển thành word",
        "chuyển đổi sang", "chuyển đổi thành",
        "convert to docx", "convert to word", "convert to file",

        # === CHO TÔI / GỬI (give me / send) ===
        "cho tôi file", "cho tôi docx", "cho tôi word", "cho tôi bản",
        "cho tôi 1 file", "cho tôi một file",
        "đưa cho tôi file", "gửi cho tôi file",
        "gửi file", "gửi tệp", "gửi bản",
        "gửi file cho tôi", "gửi docx cho tôi",
        "file cho tôi", "docx cho tôi", "word cho tôi",
        "bản word cho tôi", "bản doc cho tôi",

        # === TẠO FILE (create file) ===
        "tạo file", "tạo file docx", "tạo file word",
        "tạo tệp", "tạo bản word", "tạo bản docx",
        "làm file", "làm file docx", "làm file word",
        "generate file", "create file", "make file",

        # === RA FILE / THÀNH FILE (to file) ===
        "ra file", "thành file", "ra docx", "ra word",
        "ra file đi", "thành file đi",
        "ra bản word", "ra bản docx",

        # === SHORTHAND / CASUAL ===
        "docx đi", "word đi", "file đi", "doc đi",
        "docx cho tôi đi", "word cho tôi đi",
        "xuất đi", "tải đi", "lưu đi",
        "export đi", "download đi",

        # === IMPLICIT REQUESTS ===
        "muốn file", "muốn docx", "muốn word",
        "cần file", "cần docx", "cần word",
        "muốn tải", "muốn xuất", "muốn lưu",
        "cần tải", "cần xuất",
        "có file không", "có docx không", "có word không",
        "làm sao tải", "làm sao xuất", "làm sao lưu",

        # === ENGLISH PATTERNS ===
        "export file", "export docx", "export to docx", "export to word",
        "export as docx", "export as word", "export as pdf",
        "download file", "download docx", "download word",
        "save file", "save docx", "save as docx", "save as word",
        "give me file", "give me docx", "give me word",
        "send me file", "send me docx",
        "get file", "get docx", "get the file",
        "to docx", "to word", "to file", "as docx", "as word",

        # === COPY / SAO CHÉP ===
        "copy ra file", "sao chép ra file",
        "copy thành file", "sao ra file",
    ]
    has_direct = any(pat in question_lower for pat in direct_patterns)

    return (has_export and has_file) or has_direct


# ---------------------------------------------------------------------------
# Template-based contract serving (replaces LLM-based generation)
# ---------------------------------------------------------------------------

class ContractGenerationService:
    """Serve pre-existing legal template files from the knowledge base.

    Instead of calling an LLM to *generate* a contract from scratch, this
    service finds the best-matching template document already stored in
    Neo4j, copies (or reconstructs) the original file, registers it with
    the backend, and returns a download link.
    """

    def __init__(
        self,
        *,
        retrieval_service: RetrievalService | None = None,
        llm_client = None,
        llm_enabled: bool = True,
    ) -> None:
        self.retrieval_service = retrieval_service or RetrievalService()
        self.llm_client = llm_client or build_default_llm_client()
        self.llm_enabled = llm_enabled

    def _is_valid_template(self, doc_title: str, file_name: str, metadata: dict | None) -> bool:
        """Verify if a document in Neo4j is actually a contract template.

        It is NOT a template if:
        1. It's a user-uploaded document (unless the name contains 'mẫu' or 'template').
        2. Its title or filename contains non-template keywords (quyết định, nghị định, luật, nghị quyết, thông tư, chỉ thị, công văn, báo cáo).
        """
        title_lower = (doc_title or "").lower()
        file_lower = (file_name or "").lower()

        # Check source type if metadata is available
        source_type = (metadata or {}).get("source_type") or (metadata or {}).get("sourceType")

        # Split into tokens by common delimiters
        import re
        delimiters = r'[\s_\-\.\(\)\[\]/,]'
        title_tokens = set(re.split(delimiters, title_lower))
        file_tokens = set(re.split(delimiters, file_lower))
        all_tokens = title_tokens.union(file_tokens)

        # Common non-template legal document abbreviations and words
        non_template_tokens = {
            "qd", "qđ", "qppl", "nd", "nđ", "tt", "nq", "cv", "pl",
            "luật", "báo cáo", "tờ trình", "quyết định", "nghị định",
            "thông tư", "nghị quyết", "chỉ thị", "công văn", "pháp lệnh",
            "hiến pháp", "chủ trương", "chính sách", "ttlt", "ttlb", "ttg", 
            "qd-ttg", "qđ-ttg", "ttg-cp", "ubnd"
        }

        # If any of the tokens matches non-template legal documents, it's not a contract template
        if all_tokens.intersection(non_template_tokens):
            return False

        # General substring checks for safety
        non_template_substrings = [
            "quyết định", "nghị định", "thông tư", "luật", "nghị quyết",
            "pháp lệnh", "hiến pháp", "chỉ thị", "công văn", "báo cáo",
            "qd-ubnd", "nđ-cp", "tt-bxt", "tt-bxd", "qh15", "qh14", "qđ-ubnd",
            "ttlt", "ttlb", "ttg", "qd-ttg", "qđ-ttg", "ttg-cp", "ubnd"
        ]
        if any(sub in title_lower for sub in non_template_substrings) or any(sub in file_lower for sub in non_template_substrings):
            return False

        # If it's a user-uploaded document and doesn't explicitly mention "mẫu" or "template",
        # then it's a private case document, not a generic system template.
        if source_type == "USER_DOCUMENT":
            if not ("mẫu" in title_lower or "template" in title_lower or "mẫu" in file_lower or "template" in file_lower):
                return False

        return True

    def generate_contract(self, request: RagQueryRequest) -> RagQueryResponse:
        """Find the best matching template and return a download link."""

        # 1. Search ALL chunks (not just SYSTEM_KB) for relevant template documents.
        search_query = request.question.lower().strip()
        if any(k in search_query for k in ["thuê nhà", "thuê phòng", "thuê trọ", "thuê mặt bằng", "thuê căn hộ", "thuê văn phòng"]):
            template_search_text = "Mẫu Hợp đồng Thuê nhà ở"
        elif any(k in search_query for k in ["lao động", "làm việc", "cộng tác viên", "thử việc"]):
            template_search_text = "Mẫu Hợp đồng Lao động"
        elif "mua bán" in search_query:
            template_search_text = "Mẫu Hợp đồng Mua bán"
        else:
            template_search_text = request.question

        embedding = self.retrieval_service.embed_question(template_search_text)
        from app.models.knowledge_models import RetrievedChunk
        raw_chunks = self.retrieval_service.repository.search_chunks(
            embedding,
            top_k=max(20, request.topKKnowledgeChunks or 5),
            query_text=template_search_text,
        )
        # Convert raw RetrievedChunk objects to RagChunkHit for uniform handling
        retrieved_chunks = [
            RagChunkHit(
                citationId=f"T{i}",
                sourceType="SYSTEM_KB",
                score=float(ch.score),
                chunkText=ch.text,
                documentId=str((ch.metadata or {}).get("document_id", "")),
                knowledgeDocumentId=str((ch.metadata or {}).get("document_id", "")),
                fileName=(ch.metadata or {}).get("file_name", ""),
                title=ch.title,
                rawChunkId=ch.chunk_id,
                metadata=ch.metadata,
            )
            for i, ch in enumerate(raw_chunks, start=1)
        ]

        # 2. Check if the matched document is a valid template
        is_valid = False
        best_chunk = None
        best_doc_id = None
        source_path = None
        doc_title = "Tài liệu"

        valid_templates = []
        for chunk in retrieved_chunks:
            temp_doc_id = chunk.knowledgeDocumentId or chunk.documentId
            temp_source_path, temp_doc_title = self._get_source_document(temp_doc_id)
            if self._is_valid_template(temp_doc_title, chunk.fileName, chunk.metadata):
                valid_templates.append((chunk, temp_doc_id, temp_source_path, temp_doc_title))

        if valid_templates:
            # 1. Try to find if any matched template starts with "Mẫu Hợp đồng" (user's custom template preference)
            preferred = None
            for vt in valid_templates:
                if vt[3].startswith("Mẫu Hợp đồng"):
                    preferred = vt
                    break
            
            # 2. If not found, fall back to the highest scoring valid template
            best_match = preferred if preferred else valid_templates[0]
            
            is_valid = True
            best_chunk = best_match[0]
            best_doc_id = best_match[1]
            source_path = best_match[2]
            doc_title = best_match[3]
            
            logger.info(
                "Matched valid document template: title=%s source_path=%s score=%.4f preferred=%s",
                doc_title, source_path, best_chunk.score, preferred is not None
            )
        else:
            if retrieved_chunks:
                # If no valid template was found, log the first retrieved chunk details for debugging
                temp_doc_id = retrieved_chunks[0].knowledgeDocumentId or retrieved_chunks[0].documentId
                _, temp_doc_title = self._get_source_document(temp_doc_id)
                logger.info(
                    "No valid template found in retrieved chunks. Top chunk was: title=%s (invalid template)",
                    temp_doc_title
                )

        # 3. CASE A: Valid template exists in DB -> serve it directly
        if retrieved_chunks and is_valid:
            logger.info("Serving pre-existing template file for: %s", doc_title)
            download_links: list[str] = []
            served_path: str | None = None

            # Try to resolve path
            resolved_path = None
            if source_path and os.path.exists(source_path):
                resolved_path = source_path
            else:
                resolved_path = self._find_original_file_by_title(doc_title)
                if not resolved_path and source_path:
                    # Also try matching using the filename inside the source_path
                    filename = os.path.basename(source_path)
                    name_only = os.path.splitext(filename)[0]
                    resolved_path = self._find_original_file_by_title(name_only)

            if resolved_path and os.path.exists(resolved_path):
                # File resolved on disk — copy it to uploads
                served_path = self._copy_source_file(resolved_path, doc_title)
            else:
                # Original file is gone — reconstruct from chunks as a fallback
                served_path = self._reconstruct_docx_from_chunks(best_doc_id, doc_title)

            if served_path and os.path.exists(served_path):
                link = self._register_and_get_link(request, served_path, doc_title)
                if link:
                    download_links.append(link)

            # Build chat answer
            summary = self._build_template_summary(doc_title, retrieved_chunks)

            if download_links:
                answer = (
                    f"{summary}\n\n---\n\n"
                    "📄 **File mẫu đã sẵn sàng để tải về:**\n"
                    + "\n".join(download_links)
                )
            else:
                answer = (
                    f"{summary}\n\n"
                    "⚠️ Không thể tạo file tải về. Vui lòng liên hệ quản trị viên."
                )

            # Build citations
            citations = [
                RagCitation(
                    citationId=hit.citationId,
                    sourceType=hit.sourceType,
                    score=hit.score,
                    documentId=hit.documentId,
                    workspaceId=hit.workspaceId,
                    userId=hit.userId,
                    fileName=hit.fileName,
                    knowledgeDocumentId=hit.knowledgeDocumentId,
                    lawName=hit.lawName,
                    lawCode=hit.lawCode,
                    legalDomain=hit.legalDomain,
                    pageNumber=hit.pageNumber,
                    articleNumber=hit.articleNumber,
                    clauseNumber=hit.clauseNumber,
                    sectionTitle=hit.sectionTitle,
                )
                for hit in retrieved_chunks
            ]

            return RagQueryResponse(
                requestId=request.requestId,
                chatSessionId=request.chatSessionId,
                answer=answer,
                confidenceScore=1.0,
                shouldSuggestTicket=False,
                suggestionType="NONE",
                suggestionReason=None,
                missingInformation=None,
                riskLevel="LOW",
                legalDomain="Contract Law",
                userActionHint="CONTINUE_CHAT",
                citations=citations,
                retrievedUserChunks=0,
                retrievedKnowledgeChunks=len(retrieved_chunks),
            )

        # 4. CASE B: No valid template in DB -> dynamically draft a new one via LLM (Gemini)
        else:
            logger.info("No valid template found in DB. Dynamically drafting brand-new contract template...")
            
            # Retrieve chunks for drafting context
            user_hits = self.retrieval_service.search_user_chunks(
                request.question,
                user_id=request.userId,
                workspace_id=request.workspaceId,
                top_k=request.topKUserChunks,
            )
            from app.services.query_builder import build_legal_search_query
            legal_search_query = build_legal_search_query(request.question, user_hits)
            knowledge_hits = self.retrieval_service.search_knowledge_chunks(
                legal_search_query,
                top_k=request.topKKnowledgeChunks,
                query_text=request.question,
            )

            # Fetch available docs
            available_user_docs, available_system_docs = self._fetch_available_docs_list(request)

            # Build prompts
            from app.services.prompt_builder import build_system_prompt, build_user_prompt
            system_prompt = build_system_prompt()
            user_prompt = build_user_prompt(
                request.question,
                user_hits,
                knowledge_hits,
                chat_history=request.chatHistory,
                available_user_docs=available_user_docs,
                available_system_docs=available_system_docs,
                workspace_id=request.workspaceId,
            )

            # Explicit dynamic drafting prompt suffix
            user_prompt += (
                "\n\n⚠️ BẮT BUỘC:\n"
                "1. Bạn hãy viết đầy đủ nội dung hợp đồng hoặc văn bản được yêu cầu "
                "(chứa đầy đủ các điều khoản từ đầu tới cuối, sử dụng dấu ngoặc vuông [] làm placeholder "
                "cho thông tin cá nhân). Tuyệt đối không chỉ liệt kê cấu trúc.\n"
                "2. Tuyệt đối KHÔNG viết các câu từ chối hoặc cảnh báo giới hạn như 'Tôi không thể tạo file', 'Tôi không thể xuất file docx/word', 'Với vai trò là trợ lý AI tôi chỉ có thể cung cấp văn bản', v.v. Hãy viết thẳng vào nội dung hợp đồng vì hệ thống của chúng ta đã có module tự động chuyển tin nhắn của bạn thành file DOCX cho người dùng."
            )

            # Call LLM client
            draft_text = ""
            if self.llm_client:
                llm_result = self.llm_client.generate(system_prompt=system_prompt, user_prompt=user_prompt)
                draft_text = llm_result.answer or ""

            if not draft_text:
                return self._no_template_response(request, retrieved_chunks=[])

            # Convert markdown content to DOCX
            try:
                docx_path = self._markdown_to_docx(draft_text)
            except Exception as exc:
                logger.error("Failed to dynamically create DOCX: %s", exc)
                return self._no_template_response(request, retrieved_chunks=[])

            # Register with backend and get download link
            download_link = None
            if docx_path and os.path.exists(docx_path):
                download_link = self._register_and_get_link(
                    request, docx_path, "Hợp đồng soạn thảo bởi AI"
                )

            if download_link:
                answer = (
                    f"{draft_text}\n\n"
                    "---\n\n"
                    "📄 **File Word (.docx) của hợp đồng trên đã được tạo thành công:**\n"
                    f"{download_link}\n\n"
                    "💡 *Lưu ý: Bạn có thể tải file về để chỉnh sửa thêm.*"
                )
            else:
                answer = draft_text

            return RagQueryResponse(
                requestId=request.requestId,
                chatSessionId=request.chatSessionId,
                answer=answer,
                confidenceScore=1.0,
                shouldSuggestTicket=False,
                suggestionType="NONE",
                suggestionReason=None,
                missingInformation=None,
                riskLevel="LOW",
                legalDomain="Contract Law",
                userActionHint="CONTINUE_CHAT",
                citations=[],
                retrievedUserChunks=0,
                retrievedKnowledgeChunks=0,
            )

    # ------------------------------------------------------------------
    # Helpers — source document lookup
    # ------------------------------------------------------------------

    def _get_source_document(self, doc_id: str) -> tuple[str | None, str]:
        """Return (source_path, title) for the given Neo4j Document node."""
        if not doc_id:
            return None, "Tài liệu"
        try:
            from app.database.neo4j_client import neo4j_client

            if not neo4j_client.driver:
                neo4j_client.connect()

            records = neo4j_client.execute_query(
                "MATCH (d:Document {node_id: $doc_id}) "
                "RETURN d.source_path AS sp, d.title AS title LIMIT 1",
                {"doc_id": doc_id},
            )
            if records:
                return records[0].get("sp"), records[0].get("title") or "Tài liệu"
        except Exception as exc:
            logger.error("Failed to look up document %s: %s", doc_id, exc)
        return None, "Tài liệu"

    # ------------------------------------------------------------------
    # Helpers — file preparation
    # ------------------------------------------------------------------

    def _copy_source_file(self, source_path: str, doc_title: str) -> str | None:
        """Copy an existing source file into /app/uploads/ with a unique name."""
        try:
            ext = os.path.splitext(source_path)[1] or ".docx"
            import unicodedata
            cleaned_title = doc_title.replace('đ', 'd').replace('Đ', 'D')
            normalized = unicodedata.normalize("NFKD", cleaned_title)
            ascii_bytes = normalized.encode("ascii", "ignore")
            ascii_title = ascii_bytes.decode("ascii")
            safe_title = "".join(c if c.isalnum() or c in "._-" else "_" for c in ascii_title)[:80]
            safe_title = safe_title.replace(" ", "_")
            new_filename = f"template_{uuid.uuid4().hex}_{safe_title}{ext}"
            dest = f"/app/uploads/{new_filename}"
            shutil.copy2(source_path, dest)
            logger.info("Copied template file: %s -> %s", source_path, dest)
            return dest
        except Exception as exc:
            logger.error("Failed to copy source file %s: %s", source_path, exc)
            return None

    def _reconstruct_docx_from_chunks(self, doc_id: str, doc_title: str) -> str | None:
        """Reconstruct a DOCX file from text chunks stored in Neo4j.

        This is a fallback for when the original source file no longer exists
        on disk (e.g. /tmp files cleaned after container restart).
        """
        try:
            from app.database.neo4j_client import neo4j_client

            if not neo4j_client.driver:
                neo4j_client.connect()

            records = neo4j_client.execute_query(
                "MATCH (d:Document {node_id: $doc_id}) "
                "MATCH (c:Chunk {source_path: d.source_path}) "
                "RETURN c.text AS text, c.order AS order "
                "ORDER BY c.order ASC, c.node_id ASC",
                {"doc_id": doc_id},
            )
            if not records:
                logger.warning("No chunks found for document %s", doc_id)
                return None

            clean_paragraphs = []
            for r in records:
                text = r.get("text") or ""
                if not text.strip():
                    continue
                # Split by ' > ' to clean breadcrumb prefixes and deduplicate
                parts = text.split(" > ")
                clean_text = parts[-1].strip()
                if clean_text:
                    clean_paragraphs.append(clean_text)

            if not clean_paragraphs:
                return None

            # Build a simple DOCX
            import docx
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            document = docx.Document()
            section = document.sections[0]
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.5)
            section.right_margin = Cm(2.0)

            # Default style
            style = document.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(13)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)

            for text_line in clean_paragraphs:
                for line in text_line.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        document.add_paragraph("")
                        continue
                    para = document.add_paragraph()
                    # Center align and bold for signature blocks or headers
                    if (stripped.isupper() and len(stripped) < 100) or stripped.startswith("CỘNG HÒA") or stripped.startswith("Độc lập"):
                        run = para.add_run(stripped)
                        run.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.add_run(stripped)

            import unicodedata
            cleaned_title = doc_title.replace('đ', 'd').replace('Đ', 'D')
            normalized = unicodedata.normalize("NFKD", cleaned_title)
            ascii_bytes = normalized.encode("ascii", "ignore")
            ascii_title = ascii_bytes.decode("ascii")
            safe_title = "".join(c if c.isalnum() or c in "._-" else "_" for c in ascii_title)[:80]
            safe_title = safe_title.replace(" ", "_")
            new_filename = f"template_{uuid.uuid4().hex}_{safe_title}.docx"
            dest = f"/app/uploads/{new_filename}"
            document.save(dest)
            total_chars = sum(len(p) for p in clean_paragraphs)
            logger.info("Reconstructed DOCX from chunks: %s (%d chars)", dest, total_chars)
            return dest
        except Exception as exc:
            logger.error("Failed to reconstruct DOCX for doc %s: %s", doc_id, exc)
            return None

    # ------------------------------------------------------------------
    # Helpers — backend registration
    # ------------------------------------------------------------------

    def _register_and_get_link(self, request: RagQueryRequest, file_path: str, doc_title: str) -> str | None:
        """Register the file with the Spring Boot backend and return a download link."""
        backend_base_url = os.getenv("BACKEND_BASE_URL", "http://backend:8080")
        try:
            file_size = os.path.getsize(file_path)
            stored_filename = os.path.basename(file_path)
            ext = os.path.splitext(stored_filename)[1].lstrip(".")

            payload = {
                "workspaceId": request.workspaceId,
                "userId": request.userId,
                "originalFileName": f"{doc_title}.{ext}",
                "storedFileName": stored_filename,
                "filePath": file_path,
                "fileSize": file_size,
            }

            with httpx.Client(timeout=10.0) as client:
                resp = client.post(
                    f"{backend_base_url}/api/internal/documents/register-generated",
                    json=payload,
                )
                if resp.status_code == 200:
                    reg_id = resp.json().get("data", {}).get("documentId") or stored_filename
                    icon = "📝" if ext in ("docx", "doc") else "📥"
                    return (
                        f"{icon} [Tải về mẫu văn bản ({ext.upper()})]"
                        f"(/api/v1/workspaces/{request.workspaceId}"
                        f"/documents/{reg_id}/download)"
                    )
                else:
                    logger.error("Backend registration failed (%d): %s", resp.status_code, resp.text)
        except Exception as exc:
            logger.error("Failed to register template file: %s", exc)
        return None

    # ------------------------------------------------------------------
    # Helpers — response building
    # ------------------------------------------------------------------

    def _build_template_summary(self, doc_title: str, retrieved_chunks) -> str:
        """Build a brief summary describing the matched template."""
        preview = ""
        matched_chunk = None
        if retrieved_chunks:
            for ch in retrieved_chunks:
                ch_title = (ch.title or "").strip()
                if ch_title == doc_title.strip() or doc_title.strip() in ch_title:
                    matched_chunk = ch
                    break

        if matched_chunk:
            raw = (matched_chunk.chunkText or "").strip()
            parts = raw.split(" > ")
            clean_raw = parts[-1].strip()
            if clean_raw:
                preview = f"\n\n**Nội dung tóm tắt:**\n> {clean_raw[:500]}..."

        return (
            f"📋 **Tìm thấy mẫu văn bản phù hợp:** *{doc_title}*\n\n"
            "Hệ thống đã tìm được mẫu văn bản từ cơ sở tri thức pháp lý. "
            "Bạn có thể tải file về và chỉnh sửa theo nhu cầu thực tế."
            f"{preview}"
        )

    def _no_template_response(self, request: RagQueryRequest, *, retrieved_chunks=None) -> RagQueryResponse:
        """Return a response when no matching template is found."""
        answer = (
            "⚠️ **Không tìm thấy mẫu văn bản phù hợp**\n\n"
            "Hệ thống chưa có mẫu văn bản khớp với yêu cầu của bạn trong cơ sở tri thức. "
            "Vui lòng thử mô tả cụ thể hơn loại văn bản cần tìm, ví dụ:\n"
            "- Mẫu hợp đồng thuê nhà\n"
            "- Mẫu quyết định của UBND\n"
            "- Mẫu hợp đồng lao động\n"
        )
        return RagQueryResponse(
            requestId=request.requestId,
            chatSessionId=request.chatSessionId,
            answer=answer,
            confidenceScore=0.5,
            shouldSuggestTicket=False,
            suggestionType="ASK_MORE_INFO",
            suggestionReason="Không tìm thấy mẫu văn bản phù hợp.",
            missingInformation=None,
            riskLevel="LOW",
            legalDomain="Contract Law",
            userActionHint="CONTINUE_CHAT",
            citations=[],
            retrievedUserChunks=0,
            retrievedKnowledgeChunks=0,
        )

    # ------------------------------------------------------------------
    # Export chat content to DOCX (Robust multi-case handling)
    # ------------------------------------------------------------------

    def _fetch_available_docs_list(self, request: RagQueryRequest) -> tuple[list[str], list[str]]:
        """Fetch available document titles from Neo4j for context."""
        available_user_docs: list[str] = []
        available_system_docs: list[str] = []
        try:
            from app.database.neo4j_client import neo4j_client
            import json
            if not neo4j_client.driver:
                neo4j_client.connect()
            docs = neo4j_client.execute_query(
                "MATCH (d:Document) RETURN d.title as title, d.metadata_json as metadata_json"
            )
            user_docs: list[str] = []
            system_docs: list[str] = []
            for d in docs:
                title = d.get("title") or "Untitled"
                metadata: dict = {}
                if d.get("metadata_json"):
                    try:
                        metadata = json.loads(d["metadata_json"])
                    except Exception:
                        pass

                ws_id = metadata.get("workspace_id") or metadata.get("workspaceId")
                u_id = metadata.get("user_id") or metadata.get("userId")

                if ws_id == request.workspaceId:
                    user_docs.append(title)
                elif metadata.get("source_type") == "SYSTEM_KB" or not u_id:
                    system_docs.append(title)

            available_user_docs = sorted(set(user_docs))
            available_system_docs = sorted(set(system_docs))
        except Exception as e:
            logger.error("Failed to fetch documents list: %s", e)
        return available_user_docs, available_system_docs

    def export_chat_to_docx(self, request: RagQueryRequest) -> RagQueryResponse:
        """Export the latest assistant message from chat history as a DOCX file.

        Handles 4 distinct scenarios:
        1. **Standard Export:** Converts the last AI response (contract draft) to DOCX.
        2. **Combined Request:** User asks to draft AND export in one query (e.g. "soạn hợp đồng thuê nhà và tải docx").
        3. **Template Export:** User wants a general template.
        4. **Empty History Guidance:** If history is empty and query has no drafting context, asks user what to draft.
        """
        # --- Check if the user query contains drafting intent (e.g. "soạn hợp đồng...", "làm hợp đồng...") ---
        q_lower = request.question.lower()
        drafting_keywords = ["soạn", "tạo", "viết", "draft", "soạn thảo", "lập", "làm bản", "soạn lại"]
        contract_keywords = [
            "hợp đồng", "thoả thuận", "thuê nhà", "thuê trọ", "văn phòng", "mặt bằng",
            "lao động", "mua bán", "dịch vụ", "chuyển nhượng", "vay tiền", "ủy quyền", "contract", "agreement"
        ]
        has_drafting_intent = any(kw in q_lower for kw in drafting_keywords)
        has_contract_intent = any(kw in q_lower for kw in contract_keywords)

        is_combined_request = has_drafting_intent and has_contract_intent

        # Try to extract the last AI message from chat history
        assistant_content = self._extract_last_assistant_message(request.chatHistory)

        generated_text_for_chat = ""

        # --- CASE 1: Combined request OR Empty history but has drafting intent ---
        if is_combined_request or (not assistant_content and has_contract_intent):
            logger.info("Handling Case 2 (Combined request or empty history with drafting intent): Drafting contract first...")
            
            # Retrieve context chunks for drafting
            user_hits = self.retrieval_service.search_user_chunks(
                request.question,
                user_id=request.userId,
                workspace_id=request.workspaceId,
                top_k=request.topKUserChunks,
            )
            from app.services.query_builder import build_legal_search_query
            legal_search_query = build_legal_search_query(request.question, user_hits)
            knowledge_hits = self.retrieval_service.search_knowledge_chunks(
                legal_search_query,
                top_k=request.topKKnowledgeChunks,
                query_text=request.question,
            )
            
            # Fetch available docs list
            available_user_docs, available_system_docs = self._fetch_available_docs_list(request)
            
            # Build prompts using existing builders
            from app.services.prompt_builder import build_system_prompt, build_user_prompt
            system_prompt = build_system_prompt()
            user_prompt = build_user_prompt(
                request.question,
                user_hits,
                knowledge_hits,
                chat_history=request.chatHistory,
                available_user_docs=available_user_docs,
                available_system_docs=available_system_docs,
                workspace_id=request.workspaceId,
            )
            
            # Add explicit instruction to generate full draft
            user_prompt += (
                "\n\n⚠️ BẮT BUỘC:\n"
                "1. Bạn hãy viết đầy đủ nội dung hợp đồng hoặc văn bản được yêu cầu "
                "(chứa đầy đủ các điều khoản từ đầu tới cuối, sử dụng dấu ngoặc vuông [] làm placeholder "
                "cho thông tin cá nhân). Tuyệt đối không chỉ liệt kê cấu trúc.\n"
                "2. Tuyệt đối KHÔNG viết các câu từ chối hoặc cảnh báo giới hạn như 'Tôi không thể tạo file', 'Tôi không thể xuất file docx/word', 'Với vai trò là trợ lý AI tôi chỉ có thể cung cấp văn bản', v.v. Hãy viết thẳng vào nội dung hợp đồng vì hệ thống của chúng ta đã có module tự động chuyển tin nhắn của bạn thành file DOCX cho người dùng."
            )

            # Call LLM client to generate draft
            if self.llm_client:
                logger.info("Calling LLM client to dynamically generate contract draft...")
                llm_result = self.llm_client.generate(system_prompt=system_prompt, user_prompt=user_prompt)
                assistant_content = llm_result.answer
                generated_text_for_chat = assistant_content
            else:
                logger.warning("LLM client not configured for dynamic drafting.")
                assistant_content = None

        # --- CASE 2: No assistant content and no drafting intent (Empty History & Out of context) ---
        if not assistant_content:
            return RagQueryResponse(
                requestId=request.requestId,
                chatSessionId=request.chatSessionId,
                answer=(
                    "⚠️ **Không tìm thấy nội dung để xuất file**\n\n"
                    "Tôi chưa có nội dung hợp đồng hoặc văn bản nào trong cuộc trò chuyện này để xuất thành file Word.\n\n"
                    "**Bạn có thể:**\n"
                    "1. Yêu cầu tôi soạn thảo trước (ví dụ: *'Soạn cho tôi hợp đồng thuê nhà'*)\n"
                    "2. Hoặc yêu cầu soạn và xuất cùng lúc (ví dụ: *'Soạn hợp đồng thuê nhà và xuất file docx cho tôi'*)"
                ),
                confidenceScore=0.8,
                shouldSuggestTicket=False,
                suggestionType="ASK_MORE_INFO",
                suggestionReason="Lịch sử chat trống và không có ngữ cảnh soạn thảo.",
                missingInformation="Cần yêu cầu soạn thảo văn bản trước hoặc gộp yêu cầu.",
                riskLevel="LOW",
                legalDomain="Contract Law",
                userActionHint="CONTINUE_CHAT",
                citations=[],
                retrievedUserChunks=0,
                retrievedKnowledgeChunks=0,
            )

        # --- Convert markdown content to DOCX ---
        try:
            docx_path = self._markdown_to_docx(assistant_content)
        except Exception as exc:
            logger.error("Failed to create DOCX from chat content: %s", exc)
            return RagQueryResponse(
                requestId=request.requestId,
                chatSessionId=request.chatSessionId,
                answer="⚠️ **Lỗi hệ thống khi chuyển đổi văn bản sang định dạng DOCX.** Vui lòng thử lại sau.",
                confidenceScore=0.5,
                shouldSuggestTicket=False,
                suggestionType="NONE",
                suggestionReason=None,
                missingInformation=None,
                riskLevel="LOW",
                legalDomain="Contract Law",
                userActionHint="CONTINUE_CHAT",
                citations=[],
                retrievedUserChunks=0,
                retrievedKnowledgeChunks=0,
            )

        # --- Register with backend and get download link ---
        download_link = None
        if docx_path and os.path.exists(docx_path):
            download_link = self._register_and_get_link(
                request, docx_path, "Văn bản soạn thảo bởi AI"
            )

        # Build final chat response
        if download_link:
            if generated_text_for_chat:
                # If we dynamically generated the draft, show both the text draft AND the download link
                answer = (
                    f"{generated_text_for_chat}\n\n"
                    "---\n\n"
                    "📄 **File Word (.docx) của hợp đồng trên đã được tạo thành công:**\n"
                    f"{download_link}\n\n"
                    "💡 *Lưu ý: Bạn có thể tải file về để chỉnh sửa thêm.*"
                )
            else:
                # Standard export flow (just showing success and link)
                answer = (
                    "✅ **Đã tạo thành công file Word (.docx) cho bạn!**\n\n"
                    "Nội dung văn bản soạn thảo gần nhất đã được chuyển đổi thành file tài liệu chuẩn.\n\n"
                    "---\n\n"
                    "📄 **Tải file tại đây:**\n"
                    f"{download_link}\n\n"
                    "💡 *Lưu ý: Đây là bản nháp tham khảo. Vui lòng rà soát lại thông tin trước khi ký kết.*"
                )
        else:
            answer = (
                "⚠️ **Đã tạo file DOCX nhưng gặp sự cố kết nối với máy chủ để tạo link tải.**\n\n"
                "Vui lòng thử lại sau hoặc liên hệ quản trị viên."
            )

        return RagQueryResponse(
            requestId=request.requestId,
            chatSessionId=request.chatSessionId,
            answer=answer,
            confidenceScore=1.0,
            shouldSuggestTicket=False,
            suggestionType="NONE",
            suggestionReason=None,
            missingInformation=None,
            riskLevel="LOW",
            legalDomain="Contract Law",
            userActionHint="CONTINUE_CHAT",
            citations=[],
            retrievedUserChunks=0,
            retrievedKnowledgeChunks=0,
            llmExecuted=False,
        )

    def _extract_last_assistant_message(self, chat_history: str | None) -> str | None:
        """Extract the last assistant/AI message from the chat history string.

        The chat history format from the backend is typically a formatted string
        with role markers. We look for the last assistant message.
        """
        if not chat_history:
            return None

        # Try JSON format first: [{"role": "assistant", "content": "..."}]
        try:
            import json
            messages = json.loads(chat_history)
            if isinstance(messages, list):
                assistant_msgs = [
                    m.get("content", "") for m in messages
                    if isinstance(m, dict)
                    and m.get("role", "").lower() in ("assistant", "ai", "bot")
                    and m.get("content", "").strip()
                ]
                if assistant_msgs:
                    return assistant_msgs[-1]
        except (json.JSONDecodeError, TypeError):
            pass

        # Try text format: "Assistant: ...\nUser: ...\nAssistant: ..."
        # Split by common role markers
        role_pattern = re.compile(
            r"^\s*(?:assistant|ai|bot|trợ lý|hệ thống)\s*:\s*",
            re.IGNORECASE | re.MULTILINE,
        )
        user_pattern = re.compile(
            r"^\s*(?:user|human|người dùng)\s*:\s*",
            re.IGNORECASE | re.MULTILINE,
        )

        # Split the chat history by role markers
        parts = re.split(
            r"(?=^\s*(?:assistant|ai|bot|trợ lý|hệ thống|user|human|người dùng)\s*:)",
            chat_history,
            flags=re.IGNORECASE | re.MULTILINE,
        )

        assistant_messages = []
        for part in parts:
            part = part.strip()
            if role_pattern.match(part):
                content = role_pattern.sub("", part, count=1).strip()
                if content:
                    assistant_messages.append(content)

        if assistant_messages:
            return assistant_messages[-1]

        # Fallback: if no role markers found but there's substantial content,
        # return the whole thing (it might be a single AI response)
        cleaned = chat_history.strip()
        if len(cleaned) > 100:
            return cleaned

        return None

    def _markdown_to_docx(self, markdown_text: str) -> str:
        """Convert markdown-formatted text to a professionally styled DOCX file.

        Returns the file path of the created DOCX.
        """
        import docx
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = docx.Document()

        # Page setup — Vietnamese legal document standard
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)

        # Default style
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(13)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        lines = markdown_text.replace("\r\n", "\n").split("\n")

        for line in lines:
            stripped = line.strip()

            # Empty line → paragraph break
            if not stripped:
                document.add_paragraph("")
                continue

            # Horizontal rule
            if re.match(r"^-{3,}$", stripped):
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("─" * 50)
                run.font.color.rgb = RGBColor(180, 180, 180)
                continue

            # Heading: # ## ###
            heading_match = re.match(r"^(#{1,3})\s+(.+)", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                # Remove bold markers from heading text
                heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
                para = document.add_paragraph()
                if level == 1:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(heading_text.upper())
                    run.bold = True
                    run.font.size = Pt(16)
                elif level == 2:
                    run = para.add_run(heading_text)
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    run = para.add_run(heading_text)
                    run.bold = True
                    run.font.size = Pt(13)
                run.font.name = "Times New Roman"
                continue

            # Bullet points: - item or * item
            bullet_match = re.match(r"^[-*✅⚠️🚫📋🔴🟡🟢⛔📄💡]\s+(.*)", stripped)
            if bullet_match:
                bullet_text = bullet_match.group(1)
                para = document.add_paragraph(style="List Bullet")
                self._add_rich_text_runs(para, bullet_text)
                continue

            # Blockquote: > text
            quote_match = re.match(r"^>\s+(.*)", stripped)
            if quote_match:
                para = document.add_paragraph()
                para.paragraph_format.left_indent = Cm(1.0)
                run = para.add_run(quote_match.group(1))
                run.italic = True
                run.font.name = "Times New Roman"
                continue

            # Numbered items: 1. item, 2. item
            numbered_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
            if numbered_match:
                para = document.add_paragraph()
                self._add_rich_text_runs(para, stripped)
                continue

            # Regular paragraph with inline formatting
            para = document.add_paragraph()
            self._add_rich_text_runs(para, stripped)

        # Save
        safe_name = f"chat_export_{uuid.uuid4().hex[:12]}"
        new_filename = f"{safe_name}.docx"
        dest = f"/app/uploads/{new_filename}"
        document.save(dest)
        logger.info("Created DOCX from chat content: %s", dest)
        return dest

    def _add_rich_text_runs(self, para, text: str) -> None:
        """Add text to a paragraph with inline formatting (bold, italic, links)."""
        # Pattern: **bold**, *italic*, [link text](url)
        pattern = re.compile(r"(\*\*(.+?)\*\*)|(\*(.+?)\*)|(\[(.+?)\]\((.+?)\))")
        last_end = 0

        for match in pattern.finditer(text):
            # Add text before the match
            if match.start() > last_end:
                run = para.add_run(text[last_end:match.start()])
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)

            if match.group(1):  # **bold**
                run = para.add_run(match.group(2))
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
            elif match.group(3):  # *italic*
                run = para.add_run(match.group(4))
                run.italic = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
            elif match.group(5):  # [link](url)
                run = para.add_run(match.group(6))
                run.underline = True
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            run = para.add_run(text[last_end:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)

    def _find_original_file_by_title(self, title: str, search_dir: str = "/app/uploads") -> str | None:
        """Find the original file path on disk by searching for a matching title."""
        if not title:
            return None

        def clean(t):
            # Keep alphanumeric characters and convert to lower case for comparison
            return re.sub(r'[^a-zA-Z0-9]', '', t).lower()

        target = clean(title)
        if not target:
            return None

        for root, dirs, files in os.walk(search_dir):
            # Skip generated templates to avoid self-reference loops
            if "template_" in root:
                continue
            for f in files:
                if f.startswith("template_"):
                    continue
                name, ext = os.path.splitext(f)
                cleaned_name = clean(name)
                # STRICT MATCH ONLY!
                if cleaned_name == target:
                    resolved_path = os.path.join(root, f)
                    logger.info("Resolved original template file for '%s' -> %s", title, resolved_path)
                    return resolved_path
        return None

