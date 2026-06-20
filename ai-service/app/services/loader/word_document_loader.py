from __future__ import annotations

import io
import zipfile
from pathlib import Path

from app.models.knowledge_models import DocumentBlock, ExtractedDocument
from app.services.loader.base_loader import DocumentLoader
from app.services.loader.document_loaders import (
    is_heading_like,
    looks_like_question,
    normalize_text,
)


class WordDocumentLoader(DocumentLoader):
    supported_extensions = (".docx",)

    def _extract_embedded_image_texts(self, source_path: Path) -> list[str]:
        try:
            with zipfile.ZipFile(source_path) as archive:
                media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
                if not media_names:
                    return []

                try:
                    from PIL import Image
                except ImportError:
                    return []

                try:
                    from app.services.ocr_service import OCRService
                except Exception:
                    return []

                if not OCRService.is_available():
                    return []

                ocr_service = OCRService()
                texts: list[str] = []
                for name in media_names:
                    try:
                        with archive.open(name) as media_file:
                            image_bytes = media_file.read()
                        with Image.open(io.BytesIO(image_bytes)) as image:
                            ocr_text = ocr_service.extract_text_from_pil_image(image)
                            if ocr_text.strip():
                                texts.append(ocr_text.strip())
                    except Exception:
                        continue
                return texts
        except Exception:
            return []

    def load(self, source_path: Path) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required to load .docx files. "
                "Install the 'python-docx' package in the runtime environment."
            ) from exc

        document = Document(str(source_path))
        blocks: list[DocumentBlock] = []

        def append_block(text: str, kind: str = "paragraph", heading_level: int | None = None) -> None:
            normalized = normalize_text(text)
            if not normalized:
                return

            blocks.append(
                DocumentBlock(
                    text=normalized,
                    kind=kind,
                    order=len(blocks) + 1,
                    heading_level=heading_level,
                    metadata={"source": "docx"},
                )
            )

        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            if not text.strip():
                continue

            style_name = getattr(paragraph.style, "name", "") or ""
            kind = "paragraph"
            heading_level = None
            if is_heading_like(text) or style_name.lower().startswith("heading"):
                kind = "heading"
                heading_level = 1
            elif looks_like_question(text):
                kind = "question"

            append_block(text, kind=kind, heading_level=heading_level)

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for cell_index, cell in enumerate(row.cells, start=1):
                    cell_text = normalize_text(cell.text or "")
                    if not cell_text:
                        continue
                    append_block(
                        f"[Table {table_index} R{row_index}C{cell_index}] {cell_text}",
                        kind="table_cell",
                    )

        for image_text in self._extract_embedded_image_texts(source_path):
            append_block(image_text, kind="ocr")

        return ExtractedDocument(
            source_path=source_path,
            title=source_path.stem,
            file_type="docx",
            blocks=blocks,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "block_count": len(blocks),
            },
        )
