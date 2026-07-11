import docx
import sys

# Read docx path from arguments or default
docx_path = sys.argv[1] if len(sys.argv) > 1 else "/app/uploads/doc_cd94c166de3247c0bbf65995eab0face_Generated_Contract.docx"

doc = docx.Document(docx_path)

print("=== INSPECTING TABLES ===")
print("Number of tables:", len(doc.tables))

for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    print("Rows:", len(table.rows), "Cols:", len(table.columns))
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            print(f"  Cell({r_idx},{c_idx}) paragraphs:")
            for p in cell.paragraphs:
                print(f"    - '{p.text}' [Alignment: {p.alignment}]")

print("\n=== INSPECTING BODY PARAGRAPHS ===")
print("Number of body paragraphs:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"Paragraph {i+1}: '{p.text[:100]}' [Indent: {p.paragraph_format.first_line_indent}]")
